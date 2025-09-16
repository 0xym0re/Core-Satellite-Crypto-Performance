# Core-Satellite-Crypto-Performance — v3
# ----------------------------------------------------------------------------------------
# Nouveautés vs v2 :
# - Benchmark dédié (sélecteur) + graphique de performance relative
# - Crypto list dynamique (CoinGecko, mcap>200M) + validation Yahoo + fallback statique
# - Composition des portefeuilles affichée (ex: 60/40 ; 60/40 + X% crypto)
# - Max Drawdown remis sous la volatilité dans le tableau comparatif
# - Rapport PDF amélioré : logo non déformé (ratio), tableau plus propre, + graph "3 portefeuilles"
# - Graphs Plotly conservés (heatmap, barres, base 100) + relative vs benchmark

import streamlit as st
import yfinance as yf
import pandas as pd
import numpy as np
import io, math, requests
from datetime import timedelta
import plotly.express as px
import plotly.graph_objects as go
import sys, subprocess, os, shutil
import plotly.io as pio
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import cm
from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Image as RLImage, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import PageBreak
from PIL import Image as PILImage

# ----------------------------------------------------------------------------------------
# Utils
# ----------------------------------------------------------------------------------------
def ensure_kaleido() -> bool:
    try:
        import kaleido  # noqa: F401
        return True
    except Exception:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "kaleido"])
            import kaleido  # noqa: F401
            return True
        except Exception:
            return False

def ensure_python_docx() -> bool:
    """
    Vérifie que le module 'docx' (python-docx) est dispo.
    Tente un 'pip install python-docx' si nécessaire.
    """
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
            import docx  # noqa: F401
            return True
        except Exception:
            return False

def _to_bytes(uploaded_file):
    if not uploaded_file:
        return None
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    return uploaded_file.read()

# ----------------------------------------------------------------------------------------
# Charte graphique
# ----------------------------------------------------------------------------------------
PRIMARY = "#4E26DF"
SECONDARY = "#7CEF17"
PERF_COLORS = ["#4E26DF","#7CEF17","#35434B","#B8A8F2","#C1E5F5","#C3F793",
               "#F2CFEE","#F2F2F2","#FCD9C4","#A7C7E7","#D4C2FC","#F9F6B2","#C4FCD2"]

def _palette():
    # Cyclique, commence par PRIMARY/SECONDARY puis autres teintes de la charte
    seen = set()
    base = [PRIMARY, SECONDARY] + [c for c in PERF_COLORS if c not in (PRIMARY, SECONDARY)]
    out = []
    for c in base:
        if c not in seen:
            out.append(c)
            seen.add(c)
    return out

def _heatmap_cmap():
    # Gradient charte : PRIMARY -> gris clair -> SECONDARY
    try:
        return LinearSegmentedColormap.from_list("cs_map", [PRIMARY, "#F2F2F2", SECONDARY], N=256)
    except Exception:
        return "viridis"
        
# ----------------------------------------------------------------------------------------
# Mappings de base (fallback)
# ----------------------------------------------------------------------------------------
asset_mapping = {
    "MSCI World": "URTH",
    "Nasdaq": "^IXIC",
    "S&P 500": "^GSPC",
    "US 10Y Yield": "^TNX",     # rendement, pas un prix
    "Dollar Index": "DX-Y.NYB",
    "Gold": "GC=F",
    "iShares Bonds Agregate": "AGGG.L",
}
crypto_static = {
    "Bitcoin (BTC$)": "BTC-USD",
    "Ethereum (ETH$)": "ETH-USD",
    "Solana (SOL)": "SOL-USD",
    "Binance Coin (BNB)": "BNB-USD",
    "XRP (XRP)": "XRP-USD",
    "Cardano (ADA)": "ADA-USD",
    "Dogecoin (DOGE)": "DOGE-USD",
    "Polygon (MATIC)": "MATIC-USD",
    "TRON (TRX)": "TRX-USD",
    "Toncoin (TON)": "TON11419-USD",
    "Polkadot (DOT)": "DOT-USD",
    "Litecoin (LTC)": "LTC-USD",
    "Bitcoin Cash (BCH)": "BCH-USD",
    "Chainlink (LINK)": "LINK-USD",
    "Stellar (XLM)": "XLM-USD",
    "Monero (XMR)": "XMR-USD",
    "Avalanche (AVAX)": "AVAX-USD",
    "Aptos (APT)": "APT-USD",
    "NEAR Protocol (NEAR)": "NEAR-USD",
    "Arbitrum (ARB)": "ARB-USD",
    "Render (RNDR)": "RNDR-USD",
    "Optimism (OP)": "OP-USD",
    "Uniswap (UNI)": "UNI-USD",
    "Cosmos (ATOM)": "ATOM-USD",
    "Filecoin (FIL)": "FIL-USD",
    "Aave (AAVE)": "AAVE-USD",
    "Sui (SUI)": "SUI-USD",
    "Maker (MKR)": "MKR-USD",
    "IOTA (IOTA)": "IOTA-USD",
    "Algorand (ALGO)": "ALGO-USD",
    "VeChain (VET)": "VET-USD",
    "Injective (INJ)": "INJ-USD",
    "Celestia (TIA)": "TIA-USD",
    "Jupiter (JUP)": "JUP-USD",
    "Synthetix (SNX)": "SNX-USD",
    "The Graph (GRT)": "GRT-USD",
    "Fetch.AI (FET)": "FET-USD",
    "Hyperliquid (HYPE)": "HYPE32196-USD",
    "Bittensor (TAO)": "TAO22974-USD",
    "Shiba Inu (SHIB)": "SHIB-USD",
    "Mantle (MNT)": "MNT-USD",
    "PEPE (PEPE)": "PEPE-USD",
    "Ondo (ONDO)": "ONDO-USD",
    "Stacks (STX)": "STX-USD",
    "Chiliz (CHZ)": "CHZ-USD",
    "Raydium (RAY)": "RAY-USD",
    "dogwifhat (WIF)": "WIF-USD",
    "Theta (THETA)": "THETA-USD",
    "Tezos (XTZ)": "XTZ-USD",
    "Morpho (MORPHO)": "MORPHO-USD",
}
us_equity_mapping = {
    "Apple (AAPL)": "AAPL",
    "Microsoft (MSFT)": "MSFT",
    "Amazon (AMZN)": "AMZN",
    "NVIDIA (NVDA)": "NVDA",
    "Alphabet (GOOGL)": "GOOGL",
    "Meta (META)": "META",
    "Tesla (TSLA)": "TSLA",
}

# ----------------------------------------------------------------------------------------
# App config
# ----------------------------------------------------------------------------------------
st.set_page_config(page_title="Alphacap Digital Assets", layout="wide")
st.title("Comparaison de performances d'actifs")

# ----------------------------------------------------------------------------------------
# Portefeuilles de base
# ----------------------------------------------------------------------------------------
portfolio_allocations = {
    "Portfolio 1": {"^GSPC": 0.60, "AGGG.L": 0.40},
    "Portfolio 2": {"^GSPC": 0.57, "AGGG.L": 0.38, "GC=F": 0.05},
}
def portfolio1_label(): return "Portefeuille 1 (60/40)"
def portfolio2_label(): return "Portefeuille 2 (60/40 + 5% Or)"

# ----------------------------------------------------------------------------------------
# Helpers data
# ----------------------------------------------------------------------------------------
@st.cache_data(ttl=3600, show_spinner=False)
def download_prices(tickers, start, end):
    if isinstance(tickers, str): tickers = [tickers]
    data = yf.download(
        tickers, start=start, end=end + pd.Timedelta(days=1),
        interval="1d", auto_adjust=False, group_by="column",
        threads=True, progress=False
    )
    # MultiIndex -> Adj Close prioritaire
    if isinstance(data.columns, pd.MultiIndex):
        if "Adj Close" in data.columns.get_level_values(0):
            df = data["Adj Close"].copy()
        elif "Close" in data.columns.get_level_values(0):
            df = data["Close"].copy()
        else:
            level0 = data.columns.levels[0][0]
            df = data[level0].copy()
    else:
        if "Adj Close" in data.columns:
            df = data["Adj Close"].to_frame(name=tickers[0])
        elif "Close" in data.columns:
            df = data["Close"].to_frame(name=tickers[0])
        else:
            df = data.to_frame(name=tickers[0])
    full_idx = pd.date_range(start=start, end=end, freq="D")
    df = df.reindex(full_idx).sort_index()
    for t in tickers:
        if t not in df.columns:
            df[t] = pd.NA
    return df

def is_crypto_ticker(t, crypto_set):
    return t in crypto_set

def renormalize_weights_if_needed(prices_df, allocations):
    tickers = [t for t in allocations if t in prices_df.columns]
    if not tickers: return {}, []
    w = np.array([allocations[t] for t in tickers], dtype=float)
    if w.sum() <= 0: return {}, []
    w = w / w.sum()
    return dict(zip(tickers, w)), tickers

def detect_data_gaps(df):
    stats = []
    for col in df.columns:
        s = df[col]
        total = len(s)
        na = int(s.isna().sum())
        max_gap = 0; current = 0
        for v in s.isna().values:
            if v: current += 1; max_gap = max(max_gap, current)
            else: current = 0
        stats.append({
            "Ticker": col,
            "Nom": col,  # remplacé plus tard
            "Days": total,
            "NA": na,
            "NA_%": round(100*na/total, 2) if total > 0 else np.nan,
            "Longest_NA_Streak": max_gap
        })
    return pd.DataFrame(stats).sort_values(["NA_%","Longest_NA_Streak"], ascending=False)

# ----------------------------------------------------------------------------------------
# Rebalancing engine
# ----------------------------------------------------------------------------------------
def portfolio_returns_buy_and_hold(prices, allocations):
    alloc_norm, tickers = renormalize_weights_if_needed(prices, allocations)
    if not tickers: return pd.Series(dtype=float)
    P = prices[tickers].copy().ffill()
    base = P.iloc[0].replace(0, np.nan)
    nav = (P.divide(base) * np.array([alloc_norm[t] for t in tickers])).sum(axis=1)
    return nav.pct_change().dropna()

def portfolio_returns_with_rebalancing(prices, allocations, freq="M"):
    alloc_norm, tickers = renormalize_weights_if_needed(prices, allocations)
    if not tickers: return pd.Series(dtype=float)
    P = prices[tickers].copy().ffill()
    R = P.pct_change().dropna(how="all")
    keys = R.index.to_period("M" if freq=="M" else "Q")
    parts = []
    for _, g in R.groupby(keys):
        cols = [c for c in g.columns if c in alloc_norm]
        if not cols: continue
        w = np.array([alloc_norm[c] for c in cols], dtype=float)
        if w.sum() <= 0: continue
        w = w / w.sum()
        parts.append((g[cols] * w).sum(axis=1))
    if not parts: return pd.Series(dtype=float)
    return pd.concat(parts).sort_index()

def portfolio_daily_returns(prices, allocations, rebal_mode):
    if rebal_mode.startswith("Buy"):
        return portfolio_returns_buy_and_hold(prices, allocations)
    elif rebal_mode.startswith("Monthly"):
        return portfolio_returns_with_rebalancing(prices, allocations, "M")
    else:
        return portfolio_returns_with_rebalancing(prices, allocations, "Q")

# ----------------------------------------------------------------------------------------
# Risk metrics
# ----------------------------------------------------------------------------------------
def drawdown_stats(series):
    cum = (1 + series).cumprod()
    peak = cum.cummax()
    dd = cum/peak - 1.0
    max_dd = dd.min() if len(dd) else np.nan
    return dd, max_dd

def compute_metrics_from_returns(r, dpy=252, rf_annual=0.0,
                                 want_sortino=True, want_calmar=True,
                                 want_var=False, want_cvar=False, var_alpha=0.95):
    if r is None or len(r) == 0: return {}

    # Perf
    cum_ret = (1 + r).prod() - 1
    cagr = (1 + cum_ret)**(dpy/len(r)) - 1

    # Moyenne/vol par période -> annualisation 'classique'
    mu_ann = r.mean() * dpy
    vol_ann = r.std() * np.sqrt(dpy)

    # Sharpe/Sortino
    excess_mu = mu_ann - rf_annual
    sharpe = excess_mu/vol_ann if vol_ann and vol_ann != 0 else np.nan

    sortino = np.nan
    if want_sortino:
        downside = r.copy(); downside[downside > 0] = 0
        down_stdev_ann = downside.std() * np.sqrt(dpy)
        sortino = (excess_mu/down_stdev_ann) if down_stdev_ann and down_stdev_ann != 0 else np.nan

    # Drawdown & Calmar
    dd, max_dd = drawdown_stats(r)
    calmar = (cagr/abs(max_dd)) if want_calmar and max_dd and max_dd != 0 else np.nan

    # VaR / CVaR (historiques) à la fréquence d'échantillonnage
    var_val = cvar_val = np.nan
    if want_var or want_cvar:
        q = np.quantile(-r.dropna(), var_alpha) if len(r.dropna())>0 else np.nan
        if want_var:  var_val = q
        if want_cvar:
            tail = -r.dropna()
            tail = tail[tail >= q]
            cvar_val = tail.mean() if len(tail) > 0 else q

    return {
        "Annualized Return %": round(cagr*100, 2),
        "Cumulative Return %": round(cum_ret*100, 2),
        "Volatility %": round(vol_ann*100, 2),
        "Max Drawdown %": round(max_dd*100, 2) if pd.notna(max_dd) else np.nan,
        "Sharpe": round(sharpe, 2),
        "Sortino": round(sortino, 2) if pd.notna(sortino) else np.nan,
        "Calmar": round(calmar, 2) if pd.notna(calmar) else np.nan,
        "VaR (daily)": round(var_val*100, 2) if pd.notna(var_val) else np.nan,
        "CVaR (daily)": round(cvar_val*100, 2) if pd.notna(cvar_val) else np.nan,
    }

# ----------------------------------------------------------------------------------------
# Plotly charts
# ----------------------------------------------------------------------------------------
def plot_cumulative_lines(df_prices, names_map, title):
    df_norm = df_prices.ffill().bfill()
    df_norm = df_norm / df_norm.iloc[0] * 100
    fig = go.Figure()
    for col in df_norm.columns:
        fig.add_trace(go.Scatter(x=df_norm.index, y=df_norm[col], mode='lines',
                                 name=names_map.get(col, col)))
    fig.update_layout(title=title, xaxis_title="", yaxis_title="Base 100",
                      legend=dict(orientation="h", y=-0.2),
                      margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

def plot_perf_bars(df_prices, names_map, title):
    df = df_prices.ffill().bfill()
    perf = (df.iloc[-1]/df.iloc[0] - 1).sort_values(ascending=False)
    s = perf.rename(index=names_map).rename("Performance")
    fig = px.bar(s, text=s.apply(lambda x: f"{x*100:.2f}%"))
    fig.update_traces(textposition='outside', cliponaxis=False)
    fig.update_layout(title=title, yaxis_title="Performance", xaxis_title="",
                      uniformtext_minsize=8, uniformtext_mode='hide',
                      margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

def plot_heatmap_corr(df_prices, names_map, title):
    R = df_prices.ffill().bfill().pct_change().dropna(how="all")
    C = R.corr()
    C.index = [names_map.get(c, c) for c in C.index]
    C.columns = [names_map.get(c, c) for c in C.columns]
    fig = px.imshow(C, text_auto=".2f", aspect="auto",
                    color_continuous_scale=["#4E26DF","#a993fa","#CAE5F5","#F2F2F2","#C3F793","#7CEF17"])
    fig.update_layout(title=title, margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

def plot_relative_vs_benchmark(df_all, benchmark_ticker, names_map, title):
    df = df_all.ffill().bfill()
    if benchmark_ticker not in df.columns:
        return go.Figure()
    bench = (df[benchmark_ticker]/df[benchmark_ticker].iloc[0])
    rel = (df.divide(bench, axis=0) - 1.0) * 100  # en %
    fig = go.Figure()
    for col in [c for c in df.columns if c != benchmark_ticker]:
        fig.add_trace(go.Scatter(x=rel.index, y=rel[col], mode='lines',
                                 name=names_map.get(col, col)))
    fig.update_layout(title=title, xaxis_title="", yaxis_title="Sur/ss perf vs benchmark (%)",
                      legend=dict(orientation="h", y=-0.2),
                      margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

def plot_portfolios_cum(nav_dict, title):
    fig = go.Figure()
    for name, r in nav_dict.items():
        if r is None or r.empty: continue
        cum = (1+r).cumprod()*100
        fig.add_trace(go.Scatter(x=cum.index, y=cum, mode='lines', name=name))
    fig.update_layout(title=title, xaxis_title="", yaxis_title="Base 100",
        legend=dict(orientation="h", y=-0.2),
        margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

def fig_to_png_bytes(fig, scale=2):
    """
    Export Plotly -> PNG via Kaleido uniquement (pas besoin de Chrome).
    Installe kaleido si absent, remonte une erreur sinon.
    """
    try:
        import kaleido  # noqa: F401
    except Exception:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "kaleido"])
            import kaleido  # noqa: F401
        except Exception as e:
            raise RuntimeError(f"Installation de kaleido impossible: {e}")
    try:
        return fig.to_image(format="png", scale=scale, engine="kaleido")
    except Exception as e:
        raise RuntimeError(f"Export Plotly→PNG (kaleido) a échoué: {e}")

def build_crypto_sleeve_nav(df_prices, crypto_allocation_pairs, crypto_mapping):
    if not crypto_allocation_pairs:
        return pd.Series(dtype=float)
    total_pct = sum(p for _, p in crypto_allocation_pairs) or 0.0
    if total_pct <= 0:
        return pd.Series(dtype=float)
    tw = []
    for name, pct in crypto_allocation_pairs:
        t = crypto_mapping.get(name)
        if t in df_prices.columns and pct > 0:
            tw.append((t, pct / total_pct))
    if not tw:
        return pd.Series(dtype=float)
    P = df_prices[[t for t, _ in tw]].copy().ffill().bfill()
    R = P.pct_change().dropna()
    w = np.array([w for _, w in tw], dtype=float)
    sleeve_r = (R * w).sum(axis=1)
    nav = (1.0 + sleeve_r).cumprod()
    nav.index = R.index
    return nav

def plot_crypto_sleeve_vs_benchmark(df_all, benchmark_ticker, sleeve_nav, names_map, title):
    df = df_all.ffill().bfill()
    if sleeve_nav is None or sleeve_nav.empty or benchmark_ticker not in df.columns:
        return go.Figure()
    bench = df[benchmark_ticker].reindex(sleeve_nav.index).dropna()
    sleeve_nav = sleeve_nav.reindex(bench.index).dropna()
    if sleeve_nav.empty or bench.empty:
        return go.Figure()
    bench_norm = bench / bench.iloc[0]
    sleeve_norm = sleeve_nav / sleeve_nav.iloc[0]
    rel = (sleeve_norm / bench_norm - 1.0) * 100.0
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=rel.index, y=rel, mode='lines', name="Poche Crypto (pondérée)"))
    fig.update_layout(title=title, xaxis_title="", yaxis_title="Sur/ss perf vs benchmark (%)",
                      legend=dict(orientation="h", y=-0.2),
                      margin=dict(l=20, r=20, t=60, b=60), template="plotly_white")
    return fig

# ----------------------------------------------------------------------------------------
# Fallback Matplotlib (export des graphes sans kaleido / chrome)
# ----------------------------------------------------------------------------------------
def mpl_fig_to_png_bytes(fig, dpi=200):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()

def make_heatmap_png_mpl(df_prices, names_map, title):
    D = df_prices.ffill().bfill().pct_change().dropna(how="all")
    C = D.corr()
    C.index = [names_map.get(c, c) for c in C.index]
    C.columns = [names_map.get(c, c) for c in C.columns]
    vals = C.values
    n, m = vals.shape

    fig, ax = plt.subplots(figsize=(10, 5))
    im = ax.imshow(vals, aspect='auto', cmap=_heatmap_cmap())
    ax.set_xticks(np.arange(m))
    ax.set_xticklabels(C.columns, rotation=45, ha='right', fontsize=8)
    ax.set_yticks(np.arange(n))
    ax.set_yticklabels(C.index, fontsize=8)
    for i in range(n):
        for j in range(m):
            ax.text(j, i, f"{vals[i, j]:.2f}", ha='center', va='center', fontsize=7)
    ax.set_title(title, color="#222")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

def make_perf_bars_png_mpl(df_prices, names_map, title):
    df = df_prices.ffill().bfill()
    perf = (df.iloc[-1] / df.iloc[0] - 1).sort_values(ascending=True) * 100.0
    labels = [names_map.get(k, k) for k in perf.index.tolist()]
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.barh(labels, perf.values, color=PRIMARY)
    for i, v in enumerate(perf.values):
        ax.text(v + (0.3 if v >= 0 else -0.3), i, f"{v:.2f}%", va='center',
                ha='left' if v >= 0 else 'right', fontsize=8, color="#222")
    ax.set_title(title, color="#222")
    ax.set_xlabel("Performance (%)")
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

def make_cumulative_lines_png_mpl(df_prices, names_map, title):
    D = df_prices.ffill().bfill()
    norm = D.divide(D.iloc[0]) * 100.0
    fig, ax = plt.subplots(figsize=(10, 5))
    pal = _palette()
    for i, col in enumerate(norm.columns):
        ax.plot(norm.index, norm[col], label=names_map.get(col, col), linewidth=1.8, color=pal[i % len(pal)])
    ax.set_title(title, color="#222")
    ax.set_ylabel("Base 100")
    ax.legend(fontsize=8, ncol=3)
    fig.autofmt_xdate()
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

def make_relative_vs_benchmark_png_mpl(df_all, benchmark_ticker, names_map, title):
    df = df_all.ffill().bfill()
    if benchmark_ticker not in df.columns:
        return None
    bench = df[benchmark_ticker] / df[benchmark_ticker].iloc[0]
    rel = (df.divide(bench, axis=0) - 1.0) * 100.0
    fig, ax = plt.subplots(figsize=(10, 5))
    pal = _palette()
    idx = 0
    for col in [c for c in df.columns if c != benchmark_ticker]:
        ax.plot(rel.index, rel[col], label=names_map.get(col, col), linewidth=1.8, color=pal[idx % len(pal)])
        idx += 1
    ax.set_title(title, color="#222")
    ax.set_ylabel("Sur/ss perf vs benchmark (%)")
    ax.legend(fontsize=8, ncol=3)
    fig.autofmt_xdate()
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

def make_portfolios_cum_png_mpl(nav_dict, title):
    fig, ax = plt.subplots(figsize=(10, 5))
    pal = _palette()
    for i, (name, r) in enumerate(nav_dict.items()):
        if r is None or r.empty:
            continue
        cum = (1 + r).cumprod() * 100.0
        ax.plot(cum.index, cum.values, label=name, linewidth=1.8, color=pal[i % len(pal)])
    ax.set_title(title, color="#222")
    ax.set_ylabel("Base 100")
    ax.legend(fontsize=8, ncol=3)
    fig.autofmt_xdate()
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

def make_crypto_sleeve_vs_benchmark_png_mpl(df_all, benchmark_ticker, sleeve_nav, title):
    df = df_all.ffill().bfill()
    if sleeve_nav is None or sleeve_nav.empty or benchmark_ticker not in df.columns:
        return None

    # Aligne la poche crypto et le benchmark sur le même index
    bench = df[benchmark_ticker].reindex(sleeve_nav.index).dropna()
    s = sleeve_nav.reindex(bench.index).dropna()
    if s.empty or bench.empty:
        return None

    bench_norm = bench / bench.iloc[0]
    sleeve_norm = s / s.iloc[0]
    rel = (sleeve_norm / bench_norm - 1.0) * 100.0  # en %

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(rel.index, rel.values, label="Poche Crypto (pondérée)", linewidth=1.8, color=PRIMARY)
    ax.set_title(title, color="#222")
    ax.set_ylabel("Sur/ss perf vs benchmark (%)")
    ax.legend(fontsize=8, ncol=1)
    fig.autofmt_xdate()
    fig.tight_layout()
    return mpl_fig_to_png_bytes(fig)

# ----------------------------------------------------------------------------------------
# PDF report
# ----------------------------------------------------------------------------------------
def keep_aspect_image(file_like, target_width_cm):
    try:
        im = PILImage.open(file_like)
        w, h = im.size
        ar = w / h
        target_w = target_width_cm * cm
        target_h = target_w / ar
        file_like.seek(0)
        return RLImage(file_like, width=target_w, height=target_h)
    except Exception:
        file_like.seek(0)
        return RLImage(file_like, width=3*cm, height=3*cm)

def generate_pdf_report(company_name, logo_file, charts_dict, metrics_df, composition_lines):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=1.2*cm, bottomMargin=1.2*cm)
    elements = []

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h2 = styles["Heading2"]
    normal = styles["BodyText"]

    cell_left = ParagraphStyle("cell_left", parent=normal, fontSize=9, leading=11,
                               spaceAfter=0, wordWrap="CJK", alignment=0)
    cell_right = ParagraphStyle("cell_right", parent=normal, fontSize=9, leading=11,
                                spaceAfter=0, wordWrap="CJK", alignment=2)

    # Header
    if logo_file is not None:
        try:
            logo_file.seek(0)
            logo_buf = io.BytesIO(logo_file.read())
            logo_img = keep_aspect_image(logo_buf, target_width_cm=4.5)
            elements.append(logo_img)
            elements.append(Spacer(1, 0.2*cm))
        except Exception:
            pass
    elements.append(Paragraph(f"{company_name} — Portfolio Report", title_style))
    elements.append(Spacer(1, 0.3*cm))

    # Compositions
    if composition_lines:
        elements.append(Paragraph("Compositions des portefeuilles", h2))
        for line in composition_lines:
            elements.append(Paragraph(line, normal))
        elements.append(Spacer(1, 0.2*cm))

    # Tableau des métriques (header vert)
    if metrics_df is not None and not metrics_df.empty:
        elements.append(Paragraph("Portfolio Metrics", h2))
        df = metrics_df.copy()
        header_cells = [Paragraph("Metric", cell_left)] + \
                       [Paragraph(str(c), cell_left) for c in df.columns.tolist()]
        data = [header_cells]
        for idx, row in df.iterrows():
            row_cells = [Paragraph(str(idx), cell_left)]
            for v in row.values:
                txt = "" if pd.isna(v) else str(v)
                try:
                    _ = float(str(v).replace(",", "."))
                    row_cells.append(Paragraph(txt, cell_right))
                except Exception:
                    row_cells.append(Paragraph(txt, cell_left))
            data.append(row_cells)

        avail_w = A4[0] - doc.leftMargin - doc.rightMargin
        first_w = 0.35 * avail_w
        other_w = (avail_w - first_w) / max(1, len(df.columns))
        col_widths = [first_w] + [other_w]*len(df.columns)

        table = Table(data, repeatRows=1, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0), rl_colors.HexColor(SECONDARY)),  # vert charte
            ('TEXTCOLOR',(0,0),(-1,0), rl_colors.white),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,0),10),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
            ('ALIGN',(1,1),(-1,-1),'RIGHT'),
            ('ALIGN',(0,0),(0,-1),'LEFT'),
            ('GRID',(0,0),(-1,-1),0.3, rl_colors.HexColor("#DDDDDD")),
            ('ROWBACKGROUNDS',(0,1),(-1,-1), [rl_colors.whitesmoke, rl_colors.HexColor("#F7F7F7")]),
            ('BOTTOMPADDING',(0,0),(-1,0),6),
            ('TOPPADDING',(0,0),(-1,0),6),
            ('LEFTPADDING',(0,0),(-1,-1),4),
            ('RIGHTPADDING',(0,0),(-1,-1),4),
        ]))
        elements.append(table)

        # 👇 AJOUT : espace + saut de page si des graphiques suivent
        elements.append(Spacer(1, 0.3*cm))
        if charts_dict:              # évite une page blanche quand il n'y a pas de graph
            elements.append(PageBreak())
    
    # Graphiques
    _ = ensure_kaleido()  # pré-check soft
        # Graphiques (accepte soit des bytes PNG, soit des figures Plotly)
    for name, fig_or_png in charts_dict.items():
        try:
            if isinstance(fig_or_png, (bytes, bytearray, memoryview)):
                png_bytes = bytes(fig_or_png)
            else:
                # Dernier recours : essaye plotly->png (si kaleido dispo), sinon message dans le PDF
                png_bytes = fig_to_png_bytes(fig_or_png, scale=2)
            elements.append(Paragraph(name, h2))
            elements.append(RLImage(io.BytesIO(png_bytes), width=17*cm, height=9*cm))
            elements.append(Spacer(1, 0.3*cm))
        except Exception as e:
            elements.append(Paragraph(f"⚠️ Impossible d’exporter le graphique « {name} » : {str(e)}", normal))
            elements.append(Spacer(1, 0.2*cm))


    # Glossaire
    elements.append(PageBreak())
    elements.append(Paragraph("Glossaire des indicateurs de risque (*)", h2))
    for line in [
        "<b>Volatilité*</b> : écart-type des rendements journaliers, annualisé (base 252).",
        "<b>Max Drawdown*</b> : pire baisse (pic-creux) cumulée sur la période.",
        "<b>Sharpe*</b> : (Rendement annualisé − Taux sans risque) / Volatilité.",
        "<b>Sortino*</b> : variante du Sharpe ne pénalisant que la volatilité baissière.",
        "<b>Calmar*</b> : Rendement annualisé / |Max Drawdown|.",
        "<b>VaR (historique, daily)*</b> : perte seuil, dépassée dans (1−α) des cas.",
        "<b>CVaR (Expected Shortfall)*</b> : perte moyenne conditionnelle au-delà de la VaR.",
    ]:
        elements.append(Paragraph(line, normal))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def _rgb_from_hex(hex_str):
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _shade_cell(cell, hex_color):
    # Remplit le fond d'une cellule (header de tableau) avec une couleur hex
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
    tcPr.append(shd)
    
def generate_docx_report(company_name, logo_bytes, charts_dict, metrics_df, composition_lines):
    # Assure la dispo de python-docx
    ok = ensure_python_docx()
    if not ok:
        raise RuntimeError("Impossible d’installer python-docx (DOCX).")

    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    primary_rgb = _rgb_from_hex(PRIMARY)
    secondary_rgb = _rgb_from_hex(SECONDARY)

    doc = Document()

    # ---------- Styles globaux ----------
    # Normal (texte) -> Calibri 10 pt
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    # Astuce pour forcer Calibri partout (East Asia fallback)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    normal.font.size = Pt(10)

    # Title / Headings -> Calibri + PRIMARY
    def _tune_style(name, size_pt, bold=True, color=primary_rgb):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        st.font.size = Pt(size_pt)
        st.font.bold = bold
        st.font.color.rgb = color

    for sty, sz in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13)):
        _tune_style(sty, sz, True, secondary_rgb)

    # ---------- Titre + Logo ----------
    # Logo (au-dessus du titre, largeur 4.5 cm, ratio conservé par Word)
    if logo_bytes:
        try:
            bio = io.BytesIO(logo_bytes); bio.seek(0)
            doc.add_picture(bio, width=Cm(4.5))
        except Exception:
            pass
            
    h = doc.add_heading(f"{company_name} — Portfolio Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Logo
    if logo_bytes:
        try:
            bio = io.BytesIO(logo_bytes); bio.seek(0)
            doc.add_picture(bio, width=Cm(4.5))
        except Exception:
            pass

    # Compositions
    if composition_lines:
        doc.add_heading("Compositions des portefeuilles", level=1)
        for line in composition_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)

    # 1) Metrics en premier
    if metrics_df is not None and not metrics_df.empty:
        doc.add_heading("Portfolio Metrics", level=1)
        rows = metrics_df.shape[0] + 1
        cols = metrics_df.shape[1] + 1
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        
        # Header (fond SECONDARY + texte blanc, gras, centré)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        for j, col in enumerate(metrics_df.columns.tolist(), start=1):
            hdr_cells[j].text = str(col)
            
        for c in hdr_cells:
            _shade_cell(c, SECONDARY)
            # Applique blanc + gras + centré
            for p in c.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = _rgb_from_hex("#FFFFFF")
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)

        # Body
        def _is_num(x):
            try:
                float(str(x).replace(",", "."))
                return True
            except Exception:
                return False

        for i, (idx, row) in enumerate(metrics_df.iterrows(), start=1):
            # Colonne "Metric"
            cell0 = table.cell(i, 0)
            cell0.text = str(idx)
            for p in cell0.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)

            # Valeurs
            for j, v in enumerate(row.values, start=1):
                cell = table.cell(i, j)
                cell.text = "" if pd.isna(v) else str(v)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if _is_num(v) else WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

        # 👇 AJOUT : saut de page avant la section Graphiques
        if charts_dict:          # évite une page blanche s'il n'y a pas de graph
            doc.add_page_break()


    # 2) Graphiques ensuite
    for name, fig_or_png in charts_dict.items():
        doc.add_heading(name, level=1)
        try:
            if isinstance(fig_or_png, (bytes, bytearray, memoryview)):
                bio = io.BytesIO(bytes(fig_or_png)); bio.seek(0)
                doc.add_picture(bio, width=Cm(17))
            else:
                # last resort (peu probable, on préfère PNG bytes déjà prêts)
                _ = ensure_kaleido()
                png = fig_to_png_bytes(fig_or_png, scale=2)
                bio = io.BytesIO(png); bio.seek(0)
                doc.add_picture(bio, width=Cm(17))
        except Exception as e:
            doc.add_paragraph(f"⚠️ Impossible d’insérer le graphique : {e}")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
# ----------------------------------------------------------------------------------------
# UI : sidebar
# ----------------------------------------------------------------------------------------
with st.sidebar:
    st.header("Paramètres")
    risk_free_rate_percent = st.number_input("Taux sans risque annuel (%)", -5.0, 20.0, 0.0, 0.1)
    rebal_mode = st.selectbox("Rebalancing", ["Buy & Hold (no rebalance)", "Monthly", "Quarterly"])
    freq_mode = st.radio(
    "Horloge de calcul (fréquence d'échantillonnage)",
    ["Daily", "Weekly"],
    index=0,
    help="Daily = jours ouvrés (252) ; Hebdo = clôture vendredi (52)."
)

    risk_measures = st.multiselect("Mesures de risque à afficher",
                                   ["Sharpe","Sortino","Calmar","VaR (daily)","CVaR (daily)"],
                                   default=["Sharpe","Sortino","Calmar"])
    var_conf = st.slider("Confiance VaR/CVaR (daily)", 0.80, 0.995, 0.95, 0.005)

    st.divider()
    st.subheader("Rapport PDF")
    company_name = st.text_input("Nom société", "Alphacap Digital Assets")
    logo_file = st.file_uploader("Logo (PNG/JPG)", type=["png","jpg","jpeg"])
    include_pdf = st.checkbox("Générer un rapport PDF à l'export", value=True)

# ----------------------------------------------------------------------------------------
# Crypto list dynamique
# ----------------------------------------------------------------------------------------
@st.cache_data(ttl=24*3600, show_spinner=False)
def build_crypto_mapping_dynamic(min_mcap_usd=2e8, pages=4):
    out = {}
    session = requests.Session()
    for page in range(1, pages+1):
        url = ("https://api.coingecko.com/api/v3/coins/markets"
               f"?vs_currency=usd&order=market_cap_desc&per_page=250&page={page}")
        r = session.get(url, timeout=15)
        if r.status_code != 200: break
        arr = r.json()
        if not arr: break
        for it in arr:
            mcap = it.get("market_cap")
            if mcap is None or mcap < min_mcap_usd: continue
            name = it.get("name","").strip()
            sym = (it.get("symbol","") or "").upper()
            if not sym: continue
            y_ticker = f"{sym}-USD"
            try:
                hist = yf.Ticker(y_ticker).history(period="5d")
                if hist is None or hist.empty: continue
                out[f"{name} ({sym})"] = y_ticker
            except Exception:
                continue
    return out

st.markdown("## 💼 Composition du portefeuille crypto")
use_dynamic_crypto = st.checkbox(
    "Utiliser la liste crypto dynamique (mcap>200M)",
    value=False,
    help="Récupère la liste via CoinGecko et vérifie la dispo sur Yahoo. Fallback : liste statique."
)

try:
    crypto_mapping = build_crypto_mapping_dynamic() if use_dynamic_crypto else crypto_static
except Exception:
    st.warning("CoinGecko indisponible — fallback sur la liste crypto statique.")
    crypto_mapping = crypto_static

# Ensemble complet + map noms
full_asset_mapping = {**asset_mapping, **crypto_mapping, **us_equity_mapping}
asset_names_map = {v: k for k, v in full_asset_mapping.items()}
crypto_tickers_set = set(crypto_mapping.values())
traditional_tickers_set = set(asset_mapping.values()) | set(us_equity_mapping.values())

# ----------------------------------------------------------------------------------------
# Poche crypto
# ----------------------------------------------------------------------------------------
crypto_options = list(crypto_mapping.keys())
crypto_allocation = []
crypto_global_pct = st.number_input("% du portefeuille total alloué à l'allocation crypto", 0.0, 100.0, 5.0, 0.5)
num_crypto = st.number_input("Nombre d'actifs cryptoactifs dans la poche", 1, 15, 1, 1)

total_pct = 0.0
for i in range(num_crypto):
    cols = st.columns([3, 1])
    with cols[0]:
        selected_crypto = st.selectbox(f"Crypto {i+1}", crypto_options, key=f"crypto_{i}")
    with cols[1]:
        default_pct = 100.0 if i==0 and num_crypto==1 else 0.0
        pct = st.number_input(f"% de la crypto {i+1} dans la poche", 0.0, 100.0, default_pct, 0.1, key=f"pct_{i}")
    crypto_allocation.append((selected_crypto, pct))
    total_pct += pct

if not np.isclose(total_pct, 100.0, atol=0.01):
    st.warning(f"⚠️ La somme des pourcentages de la poche crypto est {total_pct:.2f}%. Elle doit être ≈ 100%.")
elif crypto_global_pct <= 0:
    st.warning("⚠️ Le pourcentage global alloué à la poche crypto doit être > 0.")
else:
    st.success("✅ Répartition valide du portefeuille.")
    def build_portfolio3(portfolio1_alloc, crypto_global_pct, crypto_allocation_pairs):
        portfolio3 = {}
        classic_weight = 1 - crypto_global_pct / 100.0
        for t, w in portfolio1_alloc.items():
            portfolio3[t] = portfolio3.get(t, 0.0) + w * classic_weight
        for name, pct in crypto_allocation_pairs:
            ticker = crypto_mapping[name]
            weight = (pct / 100.0) * (crypto_global_pct / 100.0)
            portfolio3[ticker] = portfolio3.get(ticker, 0.0) + weight
        return portfolio3
    portfolio_allocations["Portfolio 3"] = build_portfolio3(portfolio_allocations["Portfolio 1"], crypto_global_pct, crypto_allocation)

# ----------------------------------------------------------------------------------------
# Sélections d'actifs & période
# ----------------------------------------------------------------------------------------
available_assets = list(full_asset_mapping.keys())
benchmark_label = st.selectbox("📌 Sélectionnez le benchmark :", available_assets, index=available_assets.index("S&P 500") if "S&P 500" in available_assets else 0)
benchmark_ticker = full_asset_mapping[benchmark_label]

timeframes = {"1 semaine":"7d","1 mois":"30d","3 mois":"90d","6 mois":"180d","1 an":"365d","2 ans":"730d","3 ans":"1095d","5 ans":"1825d"}
period_label = st.selectbox("⏳ Période :", list(timeframes.keys()))

use_custom_period = st.checkbox("Période personnalisée")
c1, c2 = st.columns(2)
with c1:
    custom_start = st.date_input("Date de début", value=pd.Timestamp.today() - pd.Timedelta(days=30), disabled=not use_custom_period)
with c2:
    custom_end = st.date_input("Date de fin", value=pd.Timestamp.today() - pd.Timedelta(days=1), disabled=not use_custom_period)

# ----------------------------------------------------------------------------------------
# Comparaison d'actifs
# ----------------------------------------------------------------------------------------
st.markdown("**Liste des actifs à comparer**")
compare_assets = [a for a in available_assets]
preselect = ["Bitcoin (BTC$)","Ethereum (ETH$)","MSCI World","S&P 500","Gold"]
safe_default = [a for a in preselect if a in compare_assets]
selected_comparisons = st.multiselect("📊 Actifs à comparer :", compare_assets, default=safe_default)
compare_tickers = [full_asset_mapping[a] for a in selected_comparisons if a in full_asset_mapping]

def align_to_business_days(df):
    # prend la dernière quote de chaque jour ouvré puis ffill (jours fériés inclus)
    return df.resample("B").last().ffill()

def align_to_weekly(df, rule="W-FRI"):
    return df.resample(rule).last().ffill()


def normalize_clock(df, crypto_set, mode):
    if mode == "Daily":
        return align_to_business_days(df), 252
    elif mode == "Weekly":
        return align_to_weekly(df), 52
    else:  # fallback sécurité
        return align_to_business_days(df), 252

# ----------------------------------------------------------------------------------------
# ANALYSE
# ----------------------------------------------------------------------------------------
if st.button("🔎 Analyser"):
    try:
        tickers_graphiques = sorted(set(compare_tickers + [benchmark_ticker]))
        tickers_portefeuilles = set()
        for alloc in portfolio_allocations.values():
            tickers_portefeuilles.update(alloc.keys())
        tickers_dl = sorted(set(tickers_graphiques + list(tickers_portefeuilles)))

        # Période
        if use_custom_period:
            start_date = pd.to_datetime(custom_start); end_date = pd.to_datetime(custom_end)
        else:
            nb_days = int(timeframes[period_label].replace('d',''))
            end_date = (pd.Timestamp.today() - pd.Timedelta(days=1)).normalize()
            start_date = end_date - pd.Timedelta(days=nb_days - 1)

        # Data
        df_raw = download_prices(tickers_dl, start_date, end_date)
        df, dpy_global = normalize_clock(df_raw, crypto_tickers_set, freq_mode)

        traditional_tickers = [t for t in traditional_tickers_set if t in df.columns]
        if traditional_tickers:
            df[traditional_tickers] = df[traditional_tickers].ffill().bfill()

        # Gaps
        gaps = detect_data_gaps(df[tickers_graphiques])
        gaps["Nom"] = gaps["Ticker"].map(asset_names_map).fillna(gaps["Ticker"])
        critical = gaps[(gaps["NA_%"] > 5) | (gaps["Longest_NA_Streak"] >= 5)]
        with st.expander("🔎 Qualité des données (gaps, NA%)", expanded=len(critical)>0):
            gdisp = gaps.copy(); gdisp["NA_%"] = gdisp["NA_%"].map(lambda x: f"{x:.2f}%")
            st.dataframe(gdisp[["Nom","Ticker","Days","NA","NA_%","Longest_NA_Streak"]],
                         use_container_width=True, height=220)
            if not critical.empty:
                st.warning("Certaines séries présentent des trous significatifs (NA%>5% ou streak≥5j).")

        label_period = (f"sur {period_label.lower()}" if not use_custom_period
                        else f"du {start_date.strftime('%d/%m/%Y')} au {end_date.strftime('%d/%m/%Y')}")

        # Graphs
        df_graph = df[tickers_graphiques]
        fig_heat = plot_heatmap_corr(df_graph, asset_names_map, f"Matrice de corrélation {label_period}")
        fig_perf = plot_perf_bars(df_graph, asset_names_map, f"Performances cumulées {label_period}")
        fig_lines = plot_cumulative_lines(df_graph, asset_names_map, f"Évolution des actifs {label_period}")
        sleeve_nav = build_crypto_sleeve_nav(df, crypto_allocation, crypto_mapping)
        fig_rel = plot_crypto_sleeve_vs_benchmark(
            df, benchmark_ticker, sleeve_nav, asset_names_map,
            f"Poche crypto vs benchmark ({asset_names_map.get(benchmark_ticker, benchmark_ticker)}) {label_period}"
        )

        st.plotly_chart(fig_heat, use_container_width=True)
        st.plotly_chart(fig_perf, use_container_width=True)
        st.plotly_chart(fig_lines, use_container_width=True)
        st.plotly_chart(fig_rel, use_container_width=True)

        # Portefeuilles & métriques
        rf_annual = risk_free_rate_percent / 100.0
        port_returns = {}
        port_names_display = {}
        port_names_display["Portfolio 1"] = portfolio1_label()
        port_names_display["Portfolio 2"] = portfolio2_label()
        if "Portfolio 3" in portfolio_allocations:
            port_names_display["Portfolio 3"] = f"Portefeuille 3 (60/40 + {crypto_global_pct:.0f}% Crypto)"

        for key, alloc in portfolio_allocations.items():
            r = portfolio_daily_returns(df, alloc, rebal_mode)
            port_returns[port_names_display.get(key, key)] = r

        def want(x): return x in risk_measures
        metrics_dict = {}
        for key, alloc in portfolio_allocations.items():
            disp = port_names_display.get(key, key)
            r = port_returns.get(disp, pd.Series(dtype=float))
            m = compute_metrics_from_returns(
                r, dpy=dpy_global, rf_annual=rf_annual,
                want_sortino=want("Sortino"), want_calmar=want("Calmar"),
                want_var=want("VaR (daily)"), want_cvar=want("CVaR (daily)"), var_alpha=var_conf
            )
            metrics_dict[disp] = m
        metrics_df = pd.DataFrame(metrics_dict)

        cols_order = ["Annualized Return %","Cumulative Return %","Volatility %","Max Drawdown %","Sharpe"]
        if "Sortino" in risk_measures: cols_order.append("Sortino")
        if "Calmar" in risk_measures: cols_order.append("Calmar")
        if "VaR (daily)" in risk_measures: cols_order.append("VaR (daily)")
        if "CVaR (daily)" in risk_measures: cols_order.append("CVaR (daily)")
        metrics_df = metrics_df.reindex(index=cols_order)

        st.markdown("### Comparaison de portefeuilles")
        st.dataframe(metrics_df, use_container_width=True, height=320)
        st.caption(f"Rebalancing : **{rebal_mode}** | RF utilisé : **{risk_free_rate_percent:.2f}%** (annualisé).")

        # Compositions
        def alloc_to_text(alloc):
            items = []
            for t, w in sorted(alloc.items(), key=lambda x: -x[1]):
                if w <= 0: continue
                items.append(f"{asset_names_map.get(t, t)} {w*100:.1f}%")
            return ", ".join(items)

        comp_lines = []
        for key, alloc in portfolio_allocations.items():
            disp = port_names_display.get(key, key)
            comp_lines.append(f"<b>{disp}</b> : {alloc_to_text(alloc)}")
        st.markdown("**Compositions :**<br>" + "<br>".join(comp_lines), unsafe_allow_html=True)

        # Graph Portefeuilles
        fig_ports = plot_portfolios_cum(port_returns, f"Performance cumulée des portefeuilles ({rebal_mode})")
        st.plotly_chart(fig_ports, use_container_width=True)

        # ---------------- Export (préparer & stocker) --------------------------
        perf_pct = (df_graph.ffill().bfill()/df_graph.ffill().bfill().iloc[0]-1)*100

        # On tente plotly->PNG et, si ça échoue, on bascule AUTOMATIQUEMENT vers Matplotlib
        def try_plotly_then_mpl(name, plotly_fig, mpl_builder):
            try:
                return fig_to_png_bytes(plotly_fig, scale=2)
            except Exception:
                # Fallback pur Matplotlib (aucune dépendance)
                return mpl_builder()

        charts_for_pdf = {
            "Matrice de corrélation": try_plotly_then_mpl(
                "Matrice de corrélation", fig_heat,
                lambda: make_heatmap_png_mpl(df_graph, asset_names_map, f"Matrice de corrélation {label_period}")
            ),
            "Performances cumulées": try_plotly_then_mpl(
                "Performances cumulées", fig_perf,
                lambda: make_perf_bars_png_mpl(df_graph, asset_names_map, f"Performances cumulées {label_period}")
            ),
            "Évolution (base 100)": try_plotly_then_mpl(
                "Évolution (base 100)", fig_lines,
                lambda: make_cumulative_lines_png_mpl(df_graph, asset_names_map, f"Évolution des actifs {label_period}")
            ),
            "Perf relative vs benchmark": try_plotly_then_mpl(
                "Perf relative vs benchmark", fig_rel,
                lambda: make_crypto_sleeve_vs_benchmark_png_mpl(
                    df, benchmark_ticker, sleeve_nav,
                    f"Poche crypto vs benchmark ({asset_names_map.get(benchmark_ticker, benchmark_ticker)}) {label_period}"
                )
            ),
            "Portefeuilles (base 100)": try_plotly_then_mpl(
                "Portefeuilles (base 100)", fig_ports,
                lambda: make_portfolios_cum_png_mpl(port_returns, f"Performance cumulée des portefeuilles ({rebal_mode})")
            ),
        }

        # Nettoie les éventuels None (si un builder MPL avait échoué)
        charts_for_pdf = {k: v for k, v in charts_for_pdf.items() if v is not None}

        # Composition en texte simple pour PDF/DOCX (sans HTML)
        comp_plain = []
        for line in comp_lines:
            p = Paragraph(line.replace("<b>","").replace("</b>",""), getSampleStyleSheet()["BodyText"])
            comp_plain.append(p.getPlainText())

        st.session_state["export_payload"] = {
            "df_graph": df_graph,
            "perf_pct": perf_pct,
            "metrics_df": metrics_df,
            "gaps": gaps,
            "charts_for_pdf": charts_for_pdf,   # dict: name -> PNG bytes
            "comp_lines_plain": comp_plain,
            "company_name": company_name,
            "logo_bytes": _to_bytes(logo_file),
        }
        st.success("Résultats prêts pour export dans la section en bas de page.")

        if "US 10Y Yield" in selected_comparisons or benchmark_label == "US 10Y Yield":
            st.info("ℹ️ 'US 10Y Yield' (^TNX) est un rendement (pas un prix). Interpréter les comparaisons avec prudence.")

        # Notes risques
        st.markdown("---")
        st.markdown("### ℹ️ Notes sur les indicateurs de risque (*)")
        st.caption(
            "- **Volatilité*** : écart-type des rendements journaliers, annualisé (base 252). Plus élevé = plus instable.\n"
            "- **Max Drawdown*** : pire baisse en % entre un pic et le creux suivant.\n"
            "- **Sharpe*** : (Rendement annualisé − Taux sans risque) / Volatilité.\n"
            "- **Sortino*** : variante du Sharpe (volatilité baissière uniquement).\n"
            "- **Calmar*** : Rendement annualisé / |Max Drawdown|.\n"
            "- **VaR*** (historique, daily) : perte seuil (1−α).\n"
            "- **CVaR*** : perte moyenne conditionnelle au-delà de la VaR."
        )

    except Exception as e:
        st.error("❌ Erreur lors du chargement ou de l’analyse des données.")
        st.code(str(e))
        st.info("💡 Réessayez avec une période, un actif, ou un sous-ensemble plus restreint.")

# ================== ZONE EXPORT PERSISTANTE (hors try/except) ==================
if "export_payload" in st.session_state:
    payload = st.session_state["export_payload"]

    st.markdown("---")
    st.subheader("📥 Exporter les résultats")

    # --- Excel ---
    excel_buffer = io.BytesIO()
    try:
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            payload["df_graph"].rename(columns=asset_names_map).to_excel(writer, sheet_name="Prix")
            payload["perf_pct"].rename(columns=asset_names_map).to_excel(writer, sheet_name="Performance (%)")
            payload["metrics_df"].to_excel(writer, sheet_name="Résumé Portefeuilles")
            payload["gaps"].to_excel(writer, sheet_name="Data Gaps", index=False)
        excel_buffer.seek(0)
        st.download_button("📄 Télécharger les données & métriques (.xlsx)",
                           data=excel_buffer, file_name="donnees_completes.xlsx")
    except Exception as e_xlsx:
        st.warning(f"Export Excel indisponible : {e_xlsx}")

    # --- Rapports (PDF + DOCX) ---
    gen_both = st.button("🖨️ Générer le rapport", key="gen_both")
    if gen_both:
        logo_io = io.BytesIO(payload["logo_bytes"]) if payload["logo_bytes"] else None

        # PDF
        try:
            pdf_buf = generate_pdf_report(
                payload["company_name"],
                logo_io,
                payload["charts_for_pdf"],  # dict de PNG bytes
                payload["metrics_df"],
                composition_lines=payload["comp_lines_plain"]
            )
            st.session_state["pdf_bytes"] = pdf_buf.getvalue()
        except Exception as e:
            st.error(f"PDF: génération impossible — {e}")

        # DOCX
        try:
            docx_buf = generate_docx_report(
                payload["company_name"],
                payload["logo_bytes"],
                payload["charts_for_pdf"],  # dict de PNG bytes
                payload["metrics_df"],
                composition_lines=payload["comp_lines_plain"]
            )
            st.session_state["docx_bytes"] = docx_buf.getvalue()
        except Exception as e:
            st.error(f"Word: génération impossible — {e}")

    if "docx_bytes" in st.session_state:
        st.download_button("📝 Télécharger le rapport en Word",
                           data=st.session_state["docx_bytes"],
                           file_name="rapport_portefeuille.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if "pdf_bytes" in st.session_state:
        st.download_button("⬇️ Télécharger le rapport en PDF",
                           data=st.session_state["pdf_bytes"],
                           file_name="rapport_portefeuille.pdf",
                           mime="application/pdf")
